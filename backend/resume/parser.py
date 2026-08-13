"""
Resume parser for DOCX, PDF, and TXT formats.
Extracts structured data while preserving the original file.
"""
import re
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, field
from datetime import datetime
import docx
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table


@dataclass
class ResumeSection:
    name: str
    content: str
    level: int = 0
    order: int = 0


@dataclass
class ParsedResume:
    contact_info: Dict[str, str] = field(default_factory=dict)
    summary: str = ""
    work_history: List[Dict[str, Any]] = field(default_factory=list)
    education: List[Dict[str, Any]] = field(default_factory=list)
    certifications: List[Dict[str, Any]] = field(default_factory=list)
    skills: List[str] = field(default_factory=list)
    technical_skills: List[str] = field(default_factory=list)
    tools: List[str] = field(default_factory=list)
    projects: List[Dict[str, Any]] = field(default_factory=list)
    sections: List[ResumeSection] = field(default_factory=list)
    raw_text: str = ""
    format_info: Dict[str, Any] = field(default_factory=dict)


class ResumeParser:
    """Parse resume files into structured data."""

    SECTION_HEADERS = [
        "experience", "work experience", "employment", "professional experience",
        "education", "academic background",
        "skills", "technical skills", "core competencies", "expertise",
        "certifications", "certificates", "licenses",
        "projects", "key projects", "selected projects",
        "summary", "profile", "objective", "about",
        "achievements", "awards", "honors",
        "publications", "patents",
        "volunteer", "community",
        "languages", "interests",
    ]

    def __init__(self):
        self.format_info = {}

    def parse(self, file_path: str) -> ParsedResume:
        """Parse resume from file path."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        suffix = path.suffix.lower()
        if suffix == ".docx":
            return self._parse_docx(path)
        elif suffix == ".pdf":
            return self._parse_pdf(path)
        elif suffix in (".txt", ".md"):
            return self._parse_txt(path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    def _parse_docx(self, path: Path) -> ParsedResume:
        """Parse DOCX file preserving format info."""
        doc = docx.Document(path)

        # Extract format info
        self.format_info = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "styles": [p.style.name for p in doc.paragraphs if p.style],
            "font_info": self._extract_font_info(doc),
        }

        # Extract all text with structure
        full_text = []
        sections = []
        current_section = None
        section_order = 0

        # First pass: collect all paragraph text in order, including table cell
        # paragraphs.  Many resumes use Word tables for layout, and python-docx
        # does NOT include table-cell content in doc.paragraphs.
        all_paras: list[tuple[str, Paragraph | None]] = []
        for para in doc.paragraphs:
            all_paras.append((para.text, para))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            all_paras.append((para.text, para))

        for text, para in all_paras:
            text = text.strip()
            if not text:
                # Preserve blank lines so downstream parsers (e.g.
                # _parse_work_history) can split entries on \n\n boundaries.
                if current_section:
                    current_section.content += "\n"
                continue

            full_text.append(text)

            # Detect section headers
            is_header = self._is_section_header(text, para)
            if is_header:
                if current_section:
                    sections.append(current_section)
                current_section = ResumeSection(
                    name=text,
                    content="",
                    level=self._get_header_level(para),
                    order=section_order,
                )
                section_order += 1
            elif current_section:
                current_section.content += text + "\n"
            else:
                # Content before first section
                if not sections:
                    current_section = ResumeSection(
                        name="HEADER",
                        content=text + "\n",
                        level=0,
                        order=0,
                    )
                    section_order = 1

        if current_section:
            sections.append(current_section)

        # Merge sub-sections (level > 1) into their parent level-1 section.
        # Resumes often use "Heading 2" style for job titles / school names,
        # which the header-detection logic treats as separate sections — but
        # they contain the parent section's actual content.
        sections = self._merge_subsections(sections)

        parsed = ParsedResume(
            raw_text="\n".join(full_text),
            sections=sections,
            format_info=self.format_info,
        )

        # Parse structured content from sections
        self._parse_sections(parsed)
        return parsed

    def _extract_font_info(self, doc: DocxDocument) -> Dict[str, Any]:
        """Extract font information from document."""
        fonts = set()
        sizes = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts.add(run.font.name)
                if run.font.size:
                    sizes.add(run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size)
        return {"fonts": list(fonts), "sizes": list(sizes)}

    def _is_section_header(self, text: str, para: Paragraph) -> bool:
        """Detect if paragraph is a section header."""
        text_lower = text.lower().strip()

        # Check against known headers
        for header in self.SECTION_HEADERS:
            if text_lower == header or text_lower.startswith(header + ":") or text_lower.endswith(":" + header):
                return True

        # All-caps short text: only treat as a section header if it contains a known keyword.
        # Avoids mistaking an all-caps name (e.g. "VAISHVIK PATEL") for a section.
        if text.isupper() and len(text) < 50:
            if any(h in text_lower for h in self.SECTION_HEADERS):
                return True

        if para.style and "heading" in para.style.name.lower():
            return True

        # Bold text is only treated as a section header if it contains a known keyword.
        # This avoids splitting job titles, company names, and degree names into their own
        # sections (they're bold too, but aren't section headers).
        if para.runs and all(run.bold for run in para.runs if run.text.strip()):
            if any(h in text_lower for h in self.SECTION_HEADERS):
                return True

        return False

    def _get_header_level(self, para: Paragraph) -> int:
        """Get header level from style."""
        if para.style and para.style.name:
            if "heading 1" in para.style.name.lower():
                return 1
            elif "heading 2" in para.style.name.lower():
                return 2
            elif "heading 3" in para.style.name.lower():
                return 3
        return 1

    @staticmethod
    def _merge_subsections(sections: List[ResumeSection]) -> List[ResumeSection]:
        """Merge child sections (level > 1) into their parent level-1 section.

        Resumes commonly use Heading 2 for job titles and school names.  The
        section-detection logic treats these as separate sections, which orphans
        the actual Experience / Education content.  This step merges them back
        in so the parent section has the full content to parse.
        """
        merged: list[ResumeSection] = []
        parent: ResumeSection | None = None

        for sec in sections:
            if sec.level <= 1:
                if parent is not None:
                    merged.append(parent)
                parent = sec
            elif parent is not None:
                # Sub-section: fold its name + content into the parent.
                # Insert a blank line before the sub-section name so downstream
                # parsers (e.g. _parse_work_history) can split on \n\n boundaries.
                parent.content += f"\n\n{sec.name}\n{sec.content}"
            else:
                merged.append(sec)

        if parent is not None:
            merged.append(parent)

        return merged

    def _parse_sections(self, parsed: ParsedResume):
        """Parse structured content from sections."""
        for section in parsed.sections:
            name_lower = section.name.lower()
            content = section.content.strip()

            if any(kw in name_lower for kw in ["experience", "employment", "work"]):
                parsed.work_history.extend(self._parse_work_history(content))
            elif "education" in name_lower or "academic" in name_lower:
                parsed.education.extend(self._parse_education(content))
            elif any(kw in name_lower for kw in ["skill", "competenc", "expertise"]):
                skills = self._parse_skills(content)
                # Categorize technical vs business — no duplicates
                for skill in skills:
                    if self._is_technical_skill(skill):
                        if skill not in parsed.technical_skills:
                            parsed.technical_skills.append(skill)
                    else:
                        if skill not in parsed.skills:
                            parsed.skills.append(skill)
            elif "certif" in name_lower or "license" in name_lower:
                parsed.certifications.extend(self._parse_certifications(content))
            elif "project" in name_lower:
                parsed.projects.extend(self._parse_projects(content))
            elif any(kw in name_lower for kw in ["summary", "profile", "objective", "about"]):
                parsed.summary = content
            elif section.name == "HEADER":
                parsed.contact_info = self._parse_contact_info(content)

    def _parse_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information."""
        info = {}
        lines = text.split("\n")

        for line in lines:
            line = line.strip()
            # Email - extract just the email address
            if "@" in line and "." in line:
                email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', line)
                if email_match:
                    info["email"] = email_match.group()
            # Phone - extract just the phone number
            elif re.search(r"[\d\s\-\(\)\+]{10,}", line):
                phone_match = re.search(r'[\+]?[\d\s\-\(\)]{10,}', line)
                if phone_match:
                    info["phone"] = phone_match.group().strip()
            # LinkedIn
            elif "linkedin" in line.lower():
                linkedin_match = re.search(r'linkedin\.com/\S+', line, re.I)
                if linkedin_match:
                    info["linkedin"] = linkedin_match.group()
                else:
                    info["linkedin"] = line
            # GitHub
            elif "github" in line.lower():
                github_match = re.search(r'github\.com/\S+', line, re.I)
                if github_match:
                    info["github"] = github_match.group()
                else:
                    info["github"] = line
            # Portfolio
            elif "portfolio" in line.lower():
                portfolio_match = re.search(r'https?://\S+', line)
                if portfolio_match:
                    info["portfolio"] = portfolio_match.group()
                else:
                    info["portfolio"] = line

        # First non-empty line might be name
        for line in lines:
            line = line.strip()
            if line and not any(kw in line.lower() for kw in ["@", "http", "phone", "email", "linkedin", "github"]):
                if len(line.split()) <= 4:
                    info["name"] = line
                    break

        return info

    def _parse_work_history(self, text: str) -> List[Dict[str, Any]]:
        """Parse work history entries."""
        entries = []
        # Split by double newline or date patterns
        blocks = re.split(r"\n\s*\n", text)

        for block in blocks:
            block = block.strip()
            if not block:
                continue

            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue

            entry = {"raw": block}

            # Try to extract title, company, dates
            first_line = lines[0]
            # Pattern: Title at Company | Company - Title | Title, Company
            for pattern in [
                r"^(.+?)\s+(?:at|@)\s+(.+?)(?:\s*[|\-]\s*(.+))?$",
                r"^(.+?)\s+\|\s+(.+)$",
                # Extract company from parens: "Title (Company) Date"
                r"^(.+?)\(([^)]+)\)\s*.*?(\d{4}\s*[-–—]\s*(?:\d{4}|present|current))",
                # Comma-separated: "Title, Company, 2023-Present"
                r"^(.+?),\s*(.+?),\s*(\d{4}\s*[-–—]\s*(?:\d{4}|present|current))\s*$",
                # Comma-separated without year: "Title, Company"
                r"^(.+?),\s*([A-Z][A-Za-z .&]+)$",
            ]:
                match = re.search(pattern, first_line, re.IGNORECASE)
                if match:
                    entry["title"] = match.group(1).strip()
                    entry["company"] = match.group(2).strip()
                    if match.lastindex >= 3 and match.group(3):
                        # Only use group 3 as location if it doesn't look like a date
                        g3 = match.group(3).strip()
                        if not re.match(r"\d{4}", g3):
                            entry["location"] = g3
                    break
            # Fallback: extract company from parentheses in the first line.
            # Handles formats like "Title (Company)" where a date may include
            # month names (e.g. "Oct 2022 – June 2023") that the regex above
            # cannot parse.
            if not entry.get("title"):
                paren_m = re.search(r"\(([^)]+)\)", first_line)
                if paren_m:
                    entry["company"] = paren_m.group(1).strip()
                    entry["title"] = first_line[:paren_m.start()].strip()
            # Last resort: use the first non-bullet line as company
            if not entry.get("company") and len(lines) > 1:
                entry["title"] = entry.get("title") or lines[0]
                for ln in lines[1:]:
                    if ln and not ln.startswith(("•", "-", "·", "▪", "–", "*")) and ln[0].isupper():
                        entry["company"] = ln
                        break

            # Look for dates
            for line in lines:
                # "YYYY-YYYY" or "YYYY-Present"
                date_match = re.search(r"(\d{4})\s*[-–—]\s*(\d{4}|present|current)", line, re.I)
                if not date_match:
                    # "Month YYYY – Month YYYY" or "Month YYYY – YYYY"
                    date_match = re.search(
                        r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+"
                        r"(\d{4})\s*[-–—]\s*"
                        r"(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+)?"
                        r"(\d{4}|present|current)", line, re.I)
                if date_match:
                    entry["start_date"] = date_match.group(1)
                    entry["end_date"] = date_match.group(2)
                    break

            # Remaining lines are bullets
            bullets = []
            for line in lines[1:]:
                if line.startswith(("•", "-", "·", "▪", "–", "*")):
                    bullets.append(line.lstrip("•-·▪–* ").strip())
                elif not any(kw in line.lower() for kw in ["date", "present", "current"]):
                    bullets.append(line)
            entry["bullets"] = bullets

            entries.append(entry)

        return entries

    def _parse_education(self, text: str) -> List[Dict[str, Any]]:
        """Parse education entries."""
        entries = []
        blocks = re.split(r"\n\s*\n", text)

        for block in blocks:
            block = block.strip()
            if not block:
                continue

            entry = {"raw": block, "degree": "", "school": "", "year": "", "details": []}
            lines = [l.strip() for l in block.split("\n") if l.strip()]

            if lines:
                first_line = lines[0]
                # "School (Degree) Year" format — common in resumes
                paren_m = re.search(r"^(.+?)\((.+?)\)\s*(\d{4}.*)?$", first_line)
                if paren_m:
                    entry["school"] = paren_m.group(1).strip()
                    entry["degree"] = paren_m.group(2).strip()
                    if paren_m.group(3):
                        entry["year"] = paren_m.group(3).strip()
                # "Degree | School | Year" format
                elif "|" in first_line:
                    parts = [p.strip() for p in first_line.split("|")]
                    if len(parts) >= 2:
                        entry["degree"] = parts[0]
                        entry["school"] = parts[1]
                        if len(parts) >= 3:
                            entry["year"] = parts[2]
                    else:
                        entry["degree"] = parts[0]
                # Comma-separated: "Degree, Institution, Year" or "Institution, Degree"
                elif "," in first_line:
                    parts = [p.strip() for p in first_line.split(",")]
                    year_match = re.search(r"(\d{4})", parts[-1]) if parts else None
                    if year_match and len(parts) >= 2:
                        entry["degree"] = parts[0]
                        entry["school"] = parts[1]
                        entry["year"] = year_match.group(1)
                    elif len(parts) >= 2:
                        if parts[0][:1].upper() in ("B", "M", "P", "A", "D"):
                            entry["degree"] = parts[0]
                            entry["school"] = parts[1]
                        else:
                            entry["school"] = parts[0]
                            entry["degree"] = parts[1]
                else:
                    entry["degree"] = first_line

                for line in lines[1:]:
                    if re.search(r"\d{4}", line) and not entry["year"]:
                        entry["year"] = line
                    elif not entry["school"]:
                        entry["school"] = line
                    else:
                        entry["details"].append(line)

            entries.append(entry)

        return entries

    def _parse_skills(self, text: str) -> List[str]:
        """Parse skills list."""
        skills = []
        # Split by common delimiters (comma first — skills on separate lines
        # are still comma-separated within each line).
        for delimiter in [",", ";", "•", "·", "\n", "|", "/"]:
            if delimiter in text:
                skills = [s.strip() for s in text.split(delimiter) if s.strip()]
                break
        else:
            skills = [text.strip()] if text.strip() else []

        # Further split any item that still contains a newline (e.g.
        # "\nPython Libraries: NumPy" survived the comma split).
        flat = []
        for s in skills:
            for part in s.split("\n"):
                p = part.strip()
                if p:
                    flat.append(p)
        skills = flat

        # Remove category-prefix artifacts like "Programming Languages: SQL"
        # → extract just "SQL".
        cleaned = []
        for skill in skills:
            # Strip category prefix: "Programming Languages: SQL" → "SQL"
            if ": " in skill:
                skill = skill.split(": ", 1)[1].strip()
            skill = re.sub(r"^[\-\•\·\▪\*\s]+", "", skill)
            skill = skill.strip()
            if skill and len(skill) > 1:
                cleaned.append(skill)

        return cleaned

    def _parse_certifications(self, text: str) -> List[Dict[str, Any]]:
        """Parse certifications."""
        entries = []
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            entry = {"name": line, "raw": line}
            # Try to extract date
            date_match = re.search(r"\d{4}", line)
            if date_match:
                entry["year"] = date_match.group()
            entries.append(entry)
        return entries

    def _parse_projects(self, text: str) -> List[Dict[str, Any]]:
        """Parse projects."""
        entries = []
        blocks = re.split(r"\n\s*\n", text)

        for block in blocks:
            block = block.strip()
            if not block:
                continue

            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue

            entry = {"name": lines[0], "description": "", "technologies": [], "raw": block}
            if len(lines) > 1:
                entry["description"] = " ".join(lines[1:])
                # Extract technologies
                tech_match = re.findall(r"\b(Python|Java|JavaScript|React|SQL|AWS|Docker|Kubernetes|Git|Linux|TensorFlow|PyTorch|Pandas|NumPy|Scikit-learn|Tableau|Power BI|Excel)\b", block, re.I)
                entry["technologies"] = list(set(tech_match))

            entries.append(entry)

        return entries

    def _is_technical_skill(self, skill: str) -> bool:
        """Categorize skill as technical."""
        technical_keywords = [
            "python", "java", "javascript", "typescript", "sql", "nosql",
            "aws", "azure", "gcp", "docker", "kubernetes", "terraform",
            "git", "linux", "ci/cd", "jenkins", "github", "gitlab",
            "react", "vue", "angular", "node", "django", "flask", "fastapi",
            "pandas", "numpy", "scikit", "tensorflow", "pytorch",
            "tableau", "power bi", "looker", "excel", "vba",
            "html", "css", "rest", "graphql", "api", "microservices",
        ]
        skill_lower = skill.lower()
        return any(kw in skill_lower for kw in technical_keywords)

    def _parse_pdf(self, path: Path) -> ParsedResume:
        """Parse PDF file (requires pdfplumber or similar)."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber required for PDF parsing. Install with: pip install pdfplumber")

        with pdfplumber.open(path) as pdf:
            full_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)

        # Parse as plain text
        return self._parse_txt_content("\n".join(full_text))

    def _parse_txt(self, path: Path) -> ParsedResume:
        """Parse plain text file."""
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        return self._parse_txt_content(content)

    def _parse_txt_content(self, content: str) -> ParsedResume:
        """Parse plain text content."""
        sections = []
        current_section = None
        section_order = 0

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue

            is_header = any(
                line.lower().startswith(h) or line.lower() == h or line.lower().endswith(":" + h)
                for h in self.SECTION_HEADERS
            ) or (line.isupper() and len(line) < 50 and any(h in line.lower() for h in self.SECTION_HEADERS))

            if is_header:
                if current_section:
                    sections.append(current_section)
                current_section = ResumeSection(name=line, content="", level=1, order=section_order)
                section_order += 1
            elif current_section:
                current_section.content += line + "\n"
            else:
                if not sections:
                    current_section = ResumeSection(name="HEADER", content=line + "\n", level=0, order=0)
                    section_order = 1

        if current_section:
            sections.append(current_section)

        parsed = ParsedResume(raw_text=content, sections=sections)
        self._parse_sections(parsed)
        return parsed


def parse_resume(file_path: str) -> ParsedResume:
    """Convenience function to parse a resume."""
    parser = ResumeParser()
    return parser.parse(file_path)
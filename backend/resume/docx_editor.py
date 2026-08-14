"""
DOCX Editor for format-preserving resume customization.
Modifies content while preserving fonts, margins, spacing, styles, and structure.
"""
import copy
import re
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass

import docx
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, Cm, RGBColor

from resume.parser import ParsedResume, parse_resume, ResumeSection
from utils.logger import get_logger


logger = get_logger(__name__)


@dataclass
class BulletChange:
    """Represents a change to a bullet point."""
    section_name: str
    paragraph_index: int
    original_text: str
    new_text: str
    run_index: int = 0


@dataclass
class SummaryChange:
    """Represents a change to the summary/profile section."""
    paragraph_index: int
    original_text: str
    new_text: str


@dataclass
class SkillsChange:
    """Represents a change to skills section."""
    section_name: str
    paragraph_index: int
    original_text: str
    new_text: str


class DocxEditor:
    """
    Edit a DOCX resume while preserving all formatting.
    Works by directly manipulating runs in paragraphs to keep fonts, styles, spacing intact.
    """

    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.doc: DocxDocument = docx.Document(docx_path)
        self.original_doc = copy.deepcopy(self.doc)
        self.sections = self._identify_sections()
        self.format_info = self._extract_format_info()

    def _extract_format_info(self) -> Dict[str, Any]:
        """Extract and store all format information from the document."""
        info = {
            "page_margins": {},
            "default_font": None,
            "default_size": None,
            "styles": {},
            "paragraph_spacing": {},
        }

        # Page margins
        for section in self.doc.sections:
            info["page_margins"] = {
                "top": section.top_margin,
                "bottom": section.bottom_margin,
                "left": section.left_margin,
                "right": section.right_margin,
            }
            break  # Assume all sections same

        # Default font from Normal style
        if "Normal" in self.doc.styles:
            style = self.doc.styles["Normal"]
            if style.font.name:
                info["default_font"] = style.font.name
            if style.font.size:
                info["default_size"] = style.font.size

        # Store all styles
        for style in self.doc.styles:
            if style.type == 1:  # Paragraph style
                info["styles"][style.name] = {
                    "font_name": style.font.name,
                    "font_size": style.font.size,
                    "bold": style.font.bold,
                    "italic": style.font.italic,
                    "color": str(style.font.color.rgb) if style.font.color and style.font.color.rgb else None,
                    "space_before": style.paragraph_format.space_before,
                    "space_after": style.paragraph_format.space_after,
                    "line_spacing": style.paragraph_format.line_spacing,
                    "alignment": style.paragraph_format.alignment,
                }

        return info

    def _identify_sections(self) -> List[ResumeSection]:
        """Identify document sections for targeted editing."""
        sections = []
        current_section = None
        section_order = 0

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Detect section headers
            is_header = self._is_section_header(text, para)
            if is_header:
                if current_section:
                    sections.append(current_section)
                current_section = ResumeSection(
                    name=text,
                    content="",
                    level=self._get_header_level(para),
                    order=section_order,
                )
                current_section.paragraph_indices = [i]
                section_order += 1
            elif current_section:
                current_section.content += text + "\n"
                if not hasattr(current_section, 'paragraph_indices'):
                    current_section.paragraph_indices = []
                current_section.paragraph_indices.append(i)
            else:
                # Content before first section (header/contact info)
                if not sections:
                    current_section = ResumeSection(
                        name="HEADER",
                        content=text + "\n",
                        level=0,
                        order=0,
                    )
                    current_section.paragraph_indices = [i]
                    section_order = 1

        if current_section:
            sections.append(current_section)

        return sections

    def _is_section_header(self, text: str, para: Paragraph) -> bool:
        """Detect if paragraph is a section header."""
        text_lower = text.lower().strip()

        known_headers = [
            "experience", "work experience", "employment", "professional experience",
            "education", "academic background",
            "skills", "technical skills", "core competencies", "expertise",
            "certifications", "certificates", "licenses",
            "projects", "key projects", "selected projects",
            "summary", "profile", "objective", "about",
            "achievements", "awards", "honors",
            "publications", "patents",
            "volunteer", "community",
            "languages", "interests",
        ]

        for header in known_headers:
            if text_lower == header or text_lower.startswith(header + ":") or text_lower.endswith(":" + header):
                return True

        if text.isupper() and len(text) < 50:
            return True

        if para.style and "heading" in para.style.name.lower():
            return True

        return False

    def _get_header_level(self, para: Paragraph) -> int:
        if para.style and para.style.name:
            if "heading 1" in para.style.name.lower():
                return 1
            elif "heading 2" in para.style.name.lower():
                return 2
            elif "heading 3" in para.style.name.lower():
                return 3
        return 1

    # ==================== PUBLIC API ====================

    def replace_bullet(self, section_name: str, bullet_index: int, new_text: str) -> bool:
        """
        Replace a specific bullet point in a section while preserving formatting.
        """
        section = self._find_section(section_name)
        if not section or not hasattr(section, 'paragraph_indices'):
            return False

        bullet_paras = self._get_bullet_paragraphs(section)
        if bullet_index >= len(bullet_paras):
            return False

        para = self.doc.paragraphs[bullet_paras[bullet_index]]
        return self._replace_paragraph_text(para, new_text)

    def replace_summary(self, new_summary: str) -> bool:
        """Replace the summary/profile section content."""
        for section in self.sections:
            if any(kw in section.name.lower() for kw in ["summary", "profile", "objective", "about"]):
                if section.paragraph_indices:
                    # Skip section header paragraph — use the content paragraph
                    for idx in section.paragraph_indices:
                        para = self.doc.paragraphs[idx]
                        if not self._is_section_header(para.text.strip(), para):
                            return self._replace_paragraph_text(para, new_summary)
                    # Fallback to first paragraph if all are headers
                    para = self.doc.paragraphs[section.paragraph_indices[0]]
                    return self._replace_paragraph_text(para, new_summary)
        return False

    def replace_skills_text(self, section_name: str, new_skills_text: str) -> bool:
        """Replace skills section text."""
        section = self._find_section(section_name)
        if not section or not section.paragraph_indices:
            return False

        # Usually skills are in the first paragraph after header
        para = self.doc.paragraphs[section.paragraph_indices[0]]
        return self._replace_paragraph_text(para, new_skills_text)

    def reorder_bullets(self, section_name: str, new_order: List[int]) -> bool:
        """
        Reorder bullet points within a section.
        new_order is a list of original indices in the desired new order.
        """
        section = self._find_section(section_name)
        if not section:
            return False

        bullet_paras = self._get_bullet_paragraphs(section)
        if not bullet_paras:
            return False

        # Extract bullet texts with their formatting
        bullets_data = []
        for idx in bullet_paras:
            para = self.doc.paragraphs[idx]
            bullets_data.append({
                'text': para.text,
                'runs': [(run.text, run.bold, run.italic, run.underline, run.font.name, run.font.size, run.font.color.rgb if run.font.color and run.font.color.rgb else None) for run in para.runs],
                'style': para.style,
                'paragraph_format': copy.copy(para.paragraph_format),
            })

        # Reorder
        reordered = [bullets_data[i] for i in new_order if i < len(bullets_data)]

        # Write back
        for i, (para_idx, bullet_data) in enumerate(zip(bullet_paras, reordered)):
            para = self.doc.paragraphs[para_idx]
            self._write_formatted_paragraph(para, bullet_data)

        return True

    def add_bullet(self, section_name: str, text: str, position: int = -1) -> bool:
        """Add a new bullet point to a section."""
        section = self._find_section(section_name)
        if not section:
            return False

        bullet_paras = self._get_bullet_paragraphs(section)
        if not bullet_paras:
            return False

        # Use formatting from first bullet as template
        template_para = self.doc.paragraphs[bullet_paras[0]]
        template_runs = [(run.text, run.bold, run.italic, run.underline, run.font.name, run.font.size, run.font.color.rgb if run.font.color and run.font.color.rgb else None) for run in template_para.runs]

        # Create new paragraph
        new_para = self.doc.add_paragraph()
        new_para.style = template_para.style
        new_para.paragraph_format.space_before = template_para.paragraph_format.space_before
        new_para.paragraph_format.space_after = template_para.paragraph_format.space_after
        new_para.paragraph_format.line_spacing = template_para.paragraph_format.line_spacing
        new_para.paragraph_format.left_indent = template_para.paragraph_format.left_indent

        # Apply text with template formatting
        if template_runs:
            # Use first run's formatting
            _, bold, italic, underline, font_name, font_size, font_color = template_runs[0]
            run = new_para.add_run(text)
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if font_color:
                run.font.color.rgb = font_color
        else:
            new_para.add_run(text)

        # Move to correct position
        target_idx = bullet_paras[position] if position >= 0 else bullet_paras[-1] + 1
        self._move_paragraph(new_para, target_idx)

        return True

    def save(self, output_path: str) -> str:
        """Save the modified document."""
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output)
        logger.info(f"Saved customized resume to {output}")
        return str(output)

    # ==================== HELPER METHODS ====================

    def _find_section(self, name: str) -> Optional[ResumeSection]:
        """Find section by name (fuzzy match)."""
        name_lower = name.lower()
        for section in self.sections:
            if name_lower in section.name.lower() or section.name.lower() in name_lower:
                return section
        return None

    def _get_bullet_paragraphs(self, section: ResumeSection) -> List[int]:
        """Get paragraph indices that are bullet points in a section."""
        if not hasattr(section, 'paragraph_indices'):
            return []

        # For Experience / Work sections, ALL non-header, non-title paragraphs
        # are "bullets" — this resume format uses plain block paragraphs
        # without bullet characters or indentation.
        exp_keywords = ["experience", "work", "employment", "professional"]
        is_exp_section = any(kw in section.name.lower() for kw in exp_keywords)

        if is_exp_section:
            return [
                idx for idx in section.paragraph_indices
                if not self._is_section_header(self.doc.paragraphs[idx].text.strip(), self.doc.paragraphs[idx])
                and not self._is_job_title(self.doc.paragraphs[idx])
            ]

        bullet_indices = []
        for idx in section.paragraph_indices:
            para = self.doc.paragraphs[idx]
            if self._is_bullet(para):
                bullet_indices.append(idx)
        return bullet_indices

    def _is_bullet(self, para: Paragraph) -> bool:
        """Check if paragraph is a bullet point."""
        text = para.text.strip()
        if not text:
            return False

        # Check for bullet characters
        if text.startswith(("•", "-", "·", "▪", "–", "*", "○", "●")):
            return True

        # Check for numbering
        if re.match(r"^\d+[\.\)]\s", text):
            return True

        # Check paragraph format for left indent (common in bullets)
        if para.paragraph_format.left_indent and para.paragraph_format.left_indent > Pt(12):
            return True

        return False

    def _is_job_title(self, para: Paragraph) -> bool:
        """Check if paragraph is a job title line (not a bullet)."""
        text = para.text.strip()
        if not text:
            return True  # skip empty paragraphs

        # Job title lines typically have a company name reference
        company_patterns = [
            r'\([A-Z][A-Za-z0-9\s.&]+\)',   # matches (Company Name)
            r'–\s+[A-Z][A-Za-z0-9\s.&]+',    # matches – Company Name
        ]
        has_company_ref = any(re.search(p, text) for p in company_patterns)

        # Job titles often end with a date range e.g. "Oct 2022 – June 2023" or "2022 - Present"
        # Match any month name prefix + year + dash + year/Present
        has_date_range = bool(re.search(
            r'(?:[A-Z][a-z]{2,8}\s+)?(?:19|20)\d{2}\s*[–-]\s*'
            r'(?:Present|Current|(?:[A-Z][a-z]{2,8}\s+)?(?:19|20)\d{2})',
            text
        ))

        # Paragraph with a company reference AND date info → definitely a job title
        if has_company_ref and has_date_range:
            return True

        return False

    def _replace_paragraph_text(self, para: Paragraph, new_text: str) -> bool:
        """
        Replace paragraph text while preserving run formatting.
        Strategy: Keep first run's formatting, clear all runs, add new text with that formatting.
        """
        if not para.runs:
            # No runs, just add one
            run = para.add_run(new_text)
            return True

        # Capture formatting from first run that has text
        template_formatting = None
        for run in para.runs:
            if run.text.strip():
                template_formatting = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
                }
                break

        if not template_formatting:
            template_formatting = {
                'bold': False, 'italic': False, 'underline': False,
                'font_name': None, 'font_size': None, 'font_color': None,
            }

        # Clear all runs
        for run in para.runs:
            run.text = ""

        # Add new text with preserved formatting
        run = para.runs[0] if para.runs else para.add_run("")
        run.text = new_text
        run.bold = template_formatting['bold']
        run.italic = template_formatting['italic']
        run.underline = template_formatting['underline']
        if template_formatting['font_name']:
            run.font.name = template_formatting['font_name']
        if template_formatting['font_size']:
            run.font.size = template_formatting['font_size']
        if template_formatting['font_color']:
            run.font.color.rgb = template_formatting['font_color']

        return True

    def _write_formatted_paragraph(self, para: Paragraph, bullet_data: Dict[str, Any]):
        """Write a paragraph with pre-captured formatting."""
        # Clear existing runs
        for run in para.runs:
            run.text = ""

        runs_data = bullet_data['runs']
        if not runs_data:
            para.add_run(bullet_data.get('text', ''))
            return

        # Write each run with its formatting
        for i, (text, bold, italic, underline, font_name, font_size, font_color) in enumerate(runs_data):
            if i < len(para.runs):
                run = para.runs[i]
            else:
                run = para.add_run()

            run.text = text
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if font_color:
                run.font.color.rgb = font_color

        # Restore paragraph format
        pf = bullet_data['paragraph_format']
        para.paragraph_format.space_before = pf.space_before
        para.paragraph_format.space_after = pf.space_after
        para.paragraph_format.line_spacing = pf.line_spacing
        para.paragraph_format.left_indent = pf.left_indent
        para.paragraph_format.right_indent = pf.right_indent
        para.paragraph_format.first_line_indent = pf.first_line_indent

    def _move_paragraph(self, paragraph: Paragraph, target_index: int):
        """Move a paragraph to a new position in the document."""
        from docx.oxml import OxmlElement
        p_element = paragraph._element
        parent = p_element.getparent()
        parent.remove(p_element)

        # Insert at target position
        target_para = self.doc.paragraphs[target_index]
        target_para._element.addprevious(p_element)


def create_docx_editor(docx_path: str) -> DocxEditor:
    """Factory function to create a DocxEditor."""
    return DocxEditor(docx_path)


def customize_resume_from_plan(
    master_resume_path: str,
    output_path: str,
    customization_plan: Dict[str, Any],
) -> str:
    """
    Apply a customization plan to a master resume.
    
    Args:
        master_resume_path: Path to master resume DOCX
        output_path: Path to save customized resume
        customization_plan: Dict with keys:
            - summary_rewrite: new summary text
            - bullet_changes: list of {section, index, new_text}
            - skills_to_emphasize: list of skills to highlight
            - skills_to_deemphasize: list of skills to reduce
            - keywords_to_add: list of keywords to incorporate
            - section_reorder: list of section names in new order
    
    Returns:
        Path to saved customized resume
    """
    editor = DocxEditor(master_resume_path)

    # Apply summary rewrite
    if customization_plan.get("summary_rewrite"):
        editor.replace_summary(customization_plan["summary_rewrite"])

    # Apply bullet changes
    for change in customization_plan.get("bullet_changes", []):
        section = change.get("section", "Experience")
        index = change.get("index", 0)
        new_text = change.get("new_text", "")
        if not new_text:
            continue
        if section and new_text:
            editor.replace_bullet(section, index, new_text)

    # Apply skills changes
    if customization_plan.get("skills_to_emphasize") or customization_plan.get("skills_to_deemphasize"):
        # This would require more sophisticated skills section handling
        # For now, log that it's not fully implemented
        logger.warning("Skills emphasis/de-emphasis not fully implemented in DocxEditor")

    # Save
    return editor.save(output_path)
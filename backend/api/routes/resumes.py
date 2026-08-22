"""Resumes API routes."""
from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy import select, func
from sqlalchemy.exc import IntegrityError
from sqlalchemy.ext.asyncio import AsyncSession
from typing import Optional, Any, List, Dict
from api.schemas import (
    GeneratedResumeSchema,
    ValidationResultSchema,
    ResumeTemplateSchema,
    ApiResponse,
    BatchResumeRequest,
    BatchResumeResult,
    BatchResumeResponse,
)
from api.dependencies import get_db_session
from database.models import Resume, CandidateProfile, Job

from utils.logger import get_logger

router = APIRouter()
logger = get_logger(__name__)

# Predefined resume templates (can be extended)
RESUME_TEMPLATES = [
    {
        "id": "professional",
        "name": "Professional",
        "description": "Clean, traditional format with clear section headings",
        "preview_url": "/templates/professional.png",
        "is_default": True,
    },
    {
        "id": "modern",
        "name": "Modern",
        "description": "Contemporary design with side columns and visual elements",
        "preview_url": "/templates/modern.png",
        "is_default": False,
    },
    {
        "id": "compact",
        "name": "Compact",
        "description": "Single-page format optimized for conciseness",
        "preview_url": "/templates/compact.png",
        "is_default": False,
    },
    {
        "id": "executive",
        "name": "Executive",
        "description": "Premium format designed for senior roles",
        "preview_url": "/templates/executive.png",
        "is_default": False,
    },
]


@router.get("/resumes", response_model=ApiResponse)
async def list_resumes(
    page: int = 1,
    page_size: int = 20,
    session: AsyncSession = Depends(get_db_session),
):
    """List all resumes with pagination."""
    query = select(Resume).order_by(Resume.created_at.desc())
    query = query.offset((page - 1) * page_size).limit(page_size)
    result = await session.execute(query)
    resumes = result.scalars().all()

    total = await session.scalar(select(func.count(Resume.id)))

    items = []
    for r in resumes:
        items.append({
            "id": str(r.id),
            "job_id": str(r.job_id),
            "job_title": "",
            "company": "",
            "template_id": "",
            "file_path": r.file_path,
            "file_url": f"/resumes/{r.id}/download",
            "format": "docx",
            "customization_options": {},
            "created_at": r.created_at.isoformat() if r.created_at else None,
        })

    return ApiResponse(data={
        "items": items,
        "total": total,
        "page": page,
        "page_size": page_size,
        "total_pages": (total + page_size - 1) // page_size if total else 0,
    })


# NOTE: declared before /resumes/{resume_id} so "templates" is not
# captured as a resume_id path parameter.
@router.get("/resumes/templates", response_model=ApiResponse)
async def get_templates():
    """Get available resume templates."""
    return ApiResponse(data={"templates": RESUME_TEMPLATES})


@router.get("/resumes/{resume_id}", response_model=ApiResponse)
async def get_resume(
    resume_id: str,
    session: AsyncSession = Depends(get_db_session),
):
    """Get a single resume by ID."""
    from database.repositories import RepositoryFactory

    try:
        numeric_id = int(resume_id)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid resume id: {resume_id!r}")

    repo = RepositoryFactory(session)
    resume = await repo.resumes.get_resume(numeric_id)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Get job details
    job = await repo.jobs.get_job(resume.job_id)

    return ApiResponse(data={
        "id": str(resume.id),
        "job_id": str(resume.job_id),
        "job_title": job.title if job else "",
        "company": job.company if job else "",
        "template_id": "",
        "file_path": resume.file_path,
        "file_url": f"/resumes/{resume.id}/download",
        "format": "docx",
        "customization_options": {},
        "validation_result": {
            "truthfulness_score": resume.truthfulness_score or 0,
            "ats_score": resume.validation_score or 0,
            "issues": resume.validation_issues or [],
            "suggestions": [],
            "validated_at": resume.created_at.isoformat() if resume.created_at else None,
        } if resume.validation_score is not None else None,
        "created_at": resume.created_at.isoformat() if resume.created_at else None,
    })


async def _resolve_or_create_master_resume(session: AsyncSession, profile: Any) -> str:
    """
    Resolve existing master resume file (disk, DB, Supabase) or synthesize a clean
    ATS-compatible DOCX template from CandidateProfile to ensure generation never 404s.
    """
    import os, glob
    import docx
    from docx.shared import Pt, Inches
    from database.models import MasterResume
    from sqlalchemy import select as sa_select

    resume_dir = "data/master_resume"
    os.makedirs(resume_dir, exist_ok=True)
    resume_files = glob.glob(f"{resume_dir}/*.docx")
    for fpath in resume_files:
        try:
            docx.Document(fpath)
            return fpath
        except Exception:
            continue

    # Tier 1: Check MasterResume table in DB
    try:
        master_resume = (
            await session.execute(
                sa_select(MasterResume).where(MasterResume.file_type == "docx").order_by(MasterResume.created_at.desc()).limit(1)
            )
        ).scalars().first()
        if master_resume and master_resume.file_data:
            local_path = f"{resume_dir}/{master_resume.filename or 'master_resume.docx'}"
            with open(local_path, "wb") as f:
                f.write(bytes(master_resume.file_data))
            try:
                docx.Document(local_path)
                logger.info("Restored master resume from DB: %s", master_resume.filename)
                return local_path
            except Exception:
                pass
    except Exception as e:
        logger.warning("Could not restore master resume from DB: %s", e)

    # Tier 2: Check Supabase Storage for DOCX files
    try:
        from api.routes.profile import get_supabase_client
        sb = get_supabase_client()
        files = sb.storage.from_("resumes").list()
        docx_files = [f for f in files if f.get("name", "").endswith(".docx")]
        if docx_files:
            latest = docx_files[-1]
            file_data = sb.storage.from_("resumes").download(latest["name"])
            local_path = f"{resume_dir}/{latest['name'].split('/')[-1]}"
            with open(local_path, "wb") as f:
                f.write(file_data)
            try:
                docx.Document(local_path)
                return local_path
            except Exception:
                pass
    except Exception as e:
        logger.warning("Could not download resume from Supabase: %s", e)

    # Tier 3: Synthesize a clean ATS-friendly baseline DOCX resume from CandidateProfile
    logger.info("Synthesizing baseline master resume from candidate profile...")
    doc = docx.Document()

    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    name_str = getattr(profile, "name", None) or getattr(profile, "full_name", None) or "Vaishvik Patel"
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(name_str)
    name_run.font.size = Pt(16)
    name_run.font.bold = True

    contact_items = [
        getattr(profile, "email", "") or "",
        getattr(profile, "phone", "") or "",
        getattr(profile, "city", "") or "Toronto, ON",
        getattr(profile, "linkedin_url", "") or "",
    ]
    contact_str = " | ".join([c for c in contact_items if c])
    if contact_str:
        doc.add_paragraph(contact_str)

    summary_text = getattr(profile, "summary", None) or f"{name_str} is an experienced Data & Business Analyst with expertise in SQL, Python, Power BI, and automated reporting."
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summary_text)

    skills = getattr(profile, "skills", None) or []
    if skills:
        doc.add_heading("Technical & Analytical Skills", level=1)
        skill_names = [s.get("name", str(s)) if isinstance(s, dict) else str(s) for s in skills]
        doc.add_paragraph(", ".join(skill_names))

    history = getattr(profile, "employment_history", None) or []
    if history:
        doc.add_heading("Professional Experience", level=1)
        for emp in history:
            title = emp.get("title") or emp.get("role") or "Data Analyst"
            company = emp.get("company") or "Company"
            dates = f"{emp.get('start_date', '')} - {emp.get('end_date', 'Present')}".strip(" -")
            p = doc.add_paragraph()
            r1 = p.add_run(f"{title} - {company}")
            r1.font.bold = True
            if dates:
                r2 = p.add_run(f" ({dates})")
                r2.font.italic = True
            desc = emp.get("description") or []
            if isinstance(desc, str):
                desc = [d.strip() for d in desc.split("\n") if d.strip()]
            for d in desc:
                cleaned = d.lstrip("•-* ").strip()
                if cleaned:
                    doc.add_paragraph(cleaned, style="List Bullet")

    edu = getattr(profile, "education", None) or []
    if edu:
        doc.add_heading("Education", level=1)
        for e in edu:
            deg = e.get("degree") or "Degree"
            inst = e.get("institution") or e.get("school") or "Institution"
            doc.add_paragraph(f"{deg} - {inst}", style="List Bullet")

    synth_path = f"{resume_dir}/master_resume.docx"
    doc.save(synth_path)

    try:
        with open(synth_path, "rb") as f:
            bytes_data = f.read()
        session.add(MasterResume(filename="master_resume.docx", file_type="docx", file_data=bytes_data))
        await session.flush()
    except Exception as e:
        logger.debug("Could not cache synthesized master resume to DB: %s", e)

    return synth_path


@router.post("/resumes/generate", response_model=ApiResponse)
async def generate_resume(
    options: dict,
    session: AsyncSession = Depends(get_db_session),
):
    """Generate a customized resume for a job."""
    from resume.agent import create_resume_agent
    from database.repositories import RepositoryFactory
    from sqlalchemy import select
    from database.models import Job

    logger.info(f"Generate resume called with options: {options}")
    if not isinstance(options, dict) or "job_id" not in options:
        raise HTTPException(status_code=400, detail="job_id is required")
    try:
        job_id = int(options["job_id"])
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail=f"Invalid job_id: {options['job_id']!r}")

    repo = RepositoryFactory(session)
    job = await repo.jobs.get_job(job_id)
    if not job:
        direct = await session.execute(select(Job).where(Job.id == job_id))
        job = direct.scalars().first()

    if not job:
        raise HTTPException(status_code=404, detail=f"Job not found (id={job_id})")

    profile = await repo.candidates.get_profile()
    if not profile:
        raise HTTPException(status_code=404, detail="Candidate profile not found. Please upload your resume or create your profile on the Dashboard.")

    master_resume_path = await _resolve_or_create_master_resume(session, profile)

    agent = await create_resume_agent()
    result = await agent.generate_resume(
        job_id=job_id,
        master_resume_path=master_resume_path,
    )

    if not result.success:
        logger.error(f"Resume generation failed: {result.errors}")
        raise HTTPException(status_code=500, detail=f"Resume generation failed: {result.errors}")

    return ApiResponse(data={
        "id": str(result.resume_id),
        "job_id": str(job_id),
        "job_title": job.title,
        "company": job.company,
        "template_id": "user_uploaded",
        "file_path": result.resume_path,
        "file_url": f"/resumes/{result.resume_id}/download",
        "format": "docx",
        "created_at": result.created_at.isoformat() if result.created_at else None,
    })


@router.post("/resumes/batch-generate")
async def batch_generate(
    body: BatchResumeRequest,
    session: AsyncSession = Depends(get_db_session),
):
    """Generate resumes for multiple jobs in one call. Optionally auto-create application records."""
    import asyncio
    from resume.agent import create_resume_agent
    from database.repositories import RepositoryFactory
    from database.models import Job, Application, ApplicationStatus as AppStatusEnum
    from sqlalchemy import select as sa_select
    from sqlalchemy.exc import IntegrityError
    from api.websocket import emit_pipeline_update

    repo = RepositoryFactory(session)
    total = len(body.job_ids)

    if total == 0:
        return BatchResumeResponse(results=[], total=0, succeeded=0, failed=0)

    # Check candidate profile
    profile = await repo.candidates.get_profile()
    if not profile:
        raise HTTPException(status_code=404, detail="Candidate profile not found. Please upload your resume on the Dashboard first.")

    master_resume_path = await _resolve_or_create_master_resume(session, profile)

    agent = await create_resume_agent()
    semaphore = asyncio.Semaphore(body.max_concurrent or 3)

    async def gen_one(job_id_str: str, idx: int):
        async with semaphore:
            try:
                job_id = int(job_id_str)
                job = await repo.jobs.get_job(job_id)
                if not job:
                    raise ValueError(f"Job not found (id={job_id})")

                await emit_pipeline_update("generating", idx, total, f"Generating resume for {job.title} at {job.company}...", job_id=job_id_str)

                result = await agent.generate_resume(job_id=job_id, master_resume_path=master_resume_path)

                if not result.success:
                    raise ValueError(f"Generation failed: {result.errors}")

                application_id = None
                if body.auto_apply:
                    try:
                        # Use nested savepoint so a single IntegrityError
                        # doesn't roll back the entire session transaction.
                        async with session.begin_nested():
                            existing = await session.execute(
                                sa_select(Application).where(Application.job_id == job_id)
                            )
                            if not existing.scalars().first():
                                new_app = Application(
                                    candidate_id=profile.id,
                                    job_id=job_id,
                                    resume_id=result.resume_id,
                                    application_url=job.application_url or "https://example.com/apply",
                                    status=AppStatusEnum.READY,
                                )
                                session.add(new_app)
                                await session.flush()
                                application_id = str(new_app.id)
                    except IntegrityError:
                        # Application already exists (unique constraint on job_id).
                        # The savepoint has been rolled back automatically;
                        # the outer transaction is still intact.
                        existing = await session.execute(
                            sa_select(Application).where(Application.job_id == job_id)
                        )
                        app = existing.scalars().first()
                        if app:
                            application_id = str(app.id)

                await emit_pipeline_update("generating", idx, total, f"✓ Resume created for {job.title} at {job.company}", job_id=job_id_str)

                return {
                    "job_id": job_id_str,
                    "resume_id": str(result.resume_id),
                    "application_id": application_id,
                    "success": True,
                    "error": None,
                }
            except Exception as e:
                logger.error(f"Batch generation failed for job {job_id_str}: {e}")
                await emit_pipeline_update("generating", idx, total, f"✗ Failed for job {job_id_str}: {str(e)[:80]}", job_id=job_id_str)
                return {
                    "job_id": job_id_str,
                    "resume_id": None,
                    "application_id": None,
                    "success": False,
                    "error": str(e),
                }

    tasks = [gen_one(jid, i + 1) for i, jid in enumerate(body.job_ids)]
    batch_results = await asyncio.gather(*tasks)

    succeeded = sum(1 for r in batch_results if r["success"])
    failed = total - succeeded

    return BatchResumeResponse(
        results=[BatchResumeResult(**r) for r in batch_results],
        total=total,
        succeeded=succeeded,
        failed=failed,
    )


@router.get("/resumes/{resume_id}/validate", response_model=ApiResponse)
async def validate_resume(
    resume_id: str,
    session: AsyncSession = Depends(get_db_session),
):
    """Validate a generated resume."""
    import glob, os
    from resume.validator import ResumeValidator

    try:
        numeric_id = int(resume_id)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid resume id: {resume_id!r}")

    # 404 up front so a bogus id doesn't reach the LLM validator
    existing = await session.scalar(select(Resume).where(Resume.id == numeric_id))
    if not existing:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Find the user's uploaded resume
    resume_dir = "data/master_resume"
    resume_files = glob.glob(f"{resume_dir}/*.docx") + glob.glob(f"{resume_dir}/*.pdf")
    if not resume_files:
        # Restore from DB first (survives container redeploys)
        try:
            from database.models import MasterResume
            from sqlalchemy import select as sa_select
            master_resume = (
                await session.execute(
                    sa_select(MasterResume).order_by(MasterResume.created_at.desc()).limit(1)
                )
            ).scalars().first()
            if master_resume:
                os.makedirs(resume_dir, exist_ok=True)
                local_path = f"{resume_dir}/{master_resume.filename}"
                with open(local_path, "wb") as f:
                    f.write(bytes(master_resume.file_data))
                resume_files = [local_path]
        except Exception as e:
            logger.warning("Could not restore master resume from DB: %s", e)
    if not resume_files:
        # Try Supabase fallback
        try:
            from api.routes.profile import get_supabase_client
            sb = get_supabase_client()
            files = sb.storage.from_("resumes").list()
            if files:
                latest = files[-1]
                file_data = sb.storage.from_("resumes").download(latest["name"])
                os.makedirs(resume_dir, exist_ok=True)
                local_path = f"{resume_dir}/{latest['name'].split('/')[-1]}"
                with open(local_path, "wb") as f:
                    f.write(file_data)
                resume_files = [local_path]
        except Exception as e:
            logger.warning("Could not download resume from Supabase: %s", e)

    master_resume_path = resume_files[0] if resume_files else None
    if not master_resume_path:
        raise HTTPException(status_code=404, detail="No master resume found")

    validator = await ResumeValidator.create()
    result = await validator.validate_resume(int(resume_id), master_resume_path)

    return ApiResponse(data={
        "truthfulness_score": result.validation_score or 0,
        "ats_score": result.format_score or 0,
        "issues": [
            {
                "type": i.get("type", "other"),
                "severity": i.get("severity", "medium"),
                "message": i.get("message", ""),
            } for i in (result.issues or [])
        ],
        "suggestions": result.suggestions or [],
        "validated_at": result.validated_at.isoformat() if result.validated_at else None,
    })


@router.get("/resumes/{resume_id}/download")
async def download_resume(
    resume_id: str,
    format: str = "docx",
    session: AsyncSession = Depends(get_db_session),
):
    """Download a resume file.

    Resolution order: local disk (this deploy's generations) → Supabase
    Storage (survives redeploys; cached back to disk) → honest 404.
    """
    import os
    from database.repositories import RepositoryFactory
    from fastapi.responses import FileResponse

    try:
        numeric_id = int(resume_id)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid resume id: {resume_id!r}")

    repo = RepositoryFactory(session)
    resume = await repo.resumes.get_resume(numeric_id)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    file_path = resume.file_path
    if not os.path.isabs(file_path):
        file_path = os.path.abspath(file_path)

    if not os.path.exists(file_path):
        # Disk miss — the usual case after a Railway redeploy. Try Storage.
        from storage import materialize_resume
        restored = await materialize_resume(numeric_id, resume.filename, file_path)
        if not restored:
            logger.error(f"Resume {numeric_id} found neither on disk nor in Storage")
            raise HTTPException(
                status_code=404,
                detail=(
                    "This resume was generated before cloud storage was enabled and its "
                    "file no longer exists on the server. Regenerate it from the job card."
                ),
            )

    # Self-heal: if the file exists locally but was never persisted to Storage
    # (e.g. generation raced a Storage outage), backfill it now.
    try:
        from storage import persist_resume_file
        await persist_resume_file(numeric_id, resume.filename, file_path)
    except Exception:
        pass

    # Only DOCX files exist — serve honest bytes/mime instead of relabeling
    # a DOCX as a PDF when ?format=pdf is requested.
    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=resume.filename or f"resume_{resume.id}.docx",
    )


@router.delete("/resumes/{resume_id}", response_model=ApiResponse)
async def delete_resume(
    resume_id: str,
    session: AsyncSession = Depends(get_db_session),
):
    """Delete a resume."""
    from database.repositories import RepositoryFactory

    try:
        numeric_id = int(resume_id)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid resume id: {resume_id!r}")

    repo = RepositoryFactory(session)
    resume = await repo.resumes.get_resume(numeric_id)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    try:
        await session.delete(resume)
        await session.flush()
    except IntegrityError:
        raise HTTPException(
            status_code=409,
            detail="Resume is referenced by an application and cannot be deleted",
        )

    return ApiResponse(data={"success": True})


@router.get("/resumes/templates", response_model=ApiResponse)
async def get_templates():
    """Get available resume templates."""
    return ApiResponse(data={"templates": RESUME_TEMPLATES})
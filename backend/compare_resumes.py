"""Compare original resume vs generated resume to show changes."""
import docx

orig = docx.Document('data/master_resume/IT RESUME VAISHVIK PATEL.docx')
gen = docx.Document('data/generated_resumes/2026-08-14_Scotiabank_Data_Analyst_v01.docx')

print('=== CHANGES FOR SCOTIABANK DATA ANALYST JOB ===')
print('=' * 80)

# Compare Rogers bullets (paragraphs 14-18)
for i in [14, 15, 16, 17, 18]:
    ot = orig.paragraphs[i].text.strip()
    gt = gen.paragraphs[i].text.strip()
    if ot != gt:
        print(f'\n📝 Bullet {i-13} (Rogers) — CHANGED:')
        print(f'  ORIGINAL:  {ot}')
        print(f'  GENERATED: {gt}')
    else:
        print(f'\n→ Bullet {i-13} (Rogers): No change')

# Compare Neo Financial bullet (paragraph 20)
ot = orig.paragraphs[20].text.strip()
gt = gen.paragraphs[20].text.strip()
if ot != gt:
    print(f'\n📝 Neo Financial — CHANGED:')
    print(f'  ORIGINAL:  {ot}')
    print(f'  GENERATED: {gt}')
else:
    print(f'\n→ Neo Financial: No change')